import pandas as pd
import docx

class readers():

    def __init__(self, file_):
        self.file_ = file_
        self.supported_formats_ = { 'XL': ['xls', 'xlsx'], 'DOC': ['docx'] }
        self.excel_ds_ = dict() # key -> fname ; value -> { 'line_num': , 'row_values': }
        self.word_ds_ = { 'headings': [], 'subheadings': [], 'paragraphs':[], 'tables':[] }

    def validate(self, mode):

        extn_ = self.file_.split('.')[-1]   
        print('Checking ->', extn_, self.supported_formats_[ mode ])
        if extn_ not in self.supported_formats_[ mode ]: 
            return False

        return True

    def readXL(self):

        if self.validate( 'XL' ) is False: return None

        sheets = pd.read_excel( self.file_, sheet_name=None, engine='openpyxl')
        # Iterate over each sheet
        for sheet_name, df in sheets.items():
            # Iterate over each row with line numbers
            for line_number, row in df.iterrows():
                # Display the line number and the row data
                #print(f"Sheet: {sheet_name} Line {line_number + 1}: {row.to_dict()}")
                ll_ = self.excel_ds_[ sheet_name ] if sheet_name in self.excel_ds_ else []
                ll_.append( { 'line_num': line_number + 1, 'row_values': row.to_dict() } )
                self.excel_ds_[ sheet_name ] = ll_

        return self.excel_ds_

    def readDOC(self):

        if self.validate( 'DOC' ) is False: return None

        doc = docx.Document( self.file_ )

        for paragraph in doc.paragraphs:
            # Check the style of the paragraph to determine if it's a heading
            if paragraph.style.name.startswith('Heading 1'):
                self.word_ds_['headings'].append(paragraph.text)
            elif paragraph.style.name.startswith('Heading 2'):
                self.word_ds_['subheadings'].append(paragraph.text)
            else:
                self.word_ds_['paragraphs'].append(paragraph.text)     

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            self.word_ds_['tables'].append(table_data)

        return self.word_ds_

if __name__ == "__main__":
    import sys

    rd_ = readers( sys.argv[1] )
    #print( rd_.readXL() )
    print( rd_.readDOC() )
